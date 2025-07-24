

import os
import json
from docx import Document
from PyPDF2 import PdfReader
import frontmatter
import re
import datetime

def extract_wiki_links(content):
    """
    Extracts wiki-style links ([[...]]) from a string.

    Args:
        content (str): The string to search for links.

    Returns:
        list: A list of found wiki links.
    """
    return re.findall(r'\[\[(.*?)\]\]', content)

def parse_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error parsing PDF {file_path}: {e}")
        return None
    return {"content": text, "metadata": {}, "links": []}

def parse_json(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return {"content": json.dumps(data, indent=2), "metadata": {}, "links": []}
    except Exception as e:
        print(f"Error parsing JSON {file_path}: {e}")
        return None

def parse_docx(file_path):
    text = ""
    try:
        document = Document(file_path)
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error parsing DOCX {file_path}: {e}")
        return None
    return {"content": text, "metadata": {}, "links": []}

def parse_markdown(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            post = frontmatter.load(f)
        
        links = extract_wiki_links(post.content)
        
        # Convert any datetime objects in metadata to strings for JSON serialization
        metadata = post.metadata
        for key, value in metadata.items():
            if isinstance(value, (datetime.datetime, datetime.date)):
                metadata[key] = value.isoformat()

        return {
            "metadata": metadata,
            "content": post.content,
            "links": links
        }
    except Exception as e:
        print(f"Error parsing Markdown {file_path}: {e}")
        return None

def parse_document(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    parsed_data = None
    if file_extension == '.pdf':
        parsed_data = parse_pdf(file_path)
    elif file_extension == '.json':
        parsed_data = parse_json(file_path)
    elif file_extension == '.docx':
        parsed_data = parse_docx(file_path)
    elif file_extension == '.md':
        parsed_data = parse_markdown(file_path)
    else:
        print(f"Unsupported file type: {file_extension} for {file_path}")
    return parsed_data

if __name__ == "__main__":
    # Example Usage:
    # Create dummy files for testing
    with open("test.pdf", "w") as f:
        f.write("This is a dummy PDF file content.") # In reality, this would be binary
    with open("test.json", "w") as f:
        json.dump({"key": "value", "number": 123}, f)
    # For .docx, you'd need to create a real docx file programmatically or manually
    # from docx import Document
    # document = Document()
    # document.add_paragraph("This is a dummy DOCX file content.")
    # document.save("test.docx")
    with open("test.md", "w") as f:
        f.write("---\ntitle: Test Doc\nauthor: Gemini\n---\n# Hello\nThis is a test markdown file with a [[wiki link]].")

    print("--- Parsing test.pdf ---")
    pdf_result = parse_document("test.pdf")
    if pdf_result:
        print(pdf_result['content'][:100]) # Print first 100 chars

    print("\n--- Parsing test.json ---")
    json_result = parse_document("test.json")
    if json_result:
        print(json_result['content'])

    # print("\n--- Parsing test.docx ---")
    # docx_result = parse_document("test.docx")
    # if docx_result:
    #     print(docx_result['content'][:100])

    print("\n--- Parsing test.md ---")
    md_result = parse_document("test.md")
    if md_result:
        print("Content:", md_result['content'])
        print("Metadata:", md_result['metadata'])
        print("Links:", md_result['links'])

    # Clean up dummy files
    os.remove("test.pdf")
    os.remove("test.json")
    os.remove("test.md")
    # os.remove("test.docx")

